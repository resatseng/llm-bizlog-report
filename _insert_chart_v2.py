"""
將記憶體需求圖表插入 Word 文檔（改進版）
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def insert_chart_to_docx(doc_path, chart_path, output_path):
    """插入圖表到合適位置"""
    doc = Document(doc_path)

    # 策略：在「三種情境成本對比」表格之後插入
    target_keywords = ["三種情境", "情境對比", "基準情境", "樂觀情境", "悲觀情境"]

    best_position = None

    # 遍歷所有段落，找到表格後的位置
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 檢查是否包含關鍵詞
        if any(keyword in text for keyword in target_keywords):
            # 找到相關段落，向後找3-5個段落（跳過表格）
            if i + 5 < len(doc.paragraphs):
                # 檢查後續是否有「投資效益分析」標題
                for j in range(i+1, min(i+10, len(doc.paragraphs))):
                    if "投資效益" in doc.paragraphs[j].text or "3." in doc.paragraphs[j].text[:3]:
                        best_position = j
                        print(f"[INFO] Found insertion point before section: {doc.paragraphs[j].text[:50]}")
                        break

    if best_position is None:
        # 備選：在「成本分析」章節結束處
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading') and "3." in para.text[:3]:
                best_position = i
                print(f"[INFO] Using fallback position before: {para.text[:50]}")
                break

    if best_position is None:
        print("[WARNING] Could not find ideal position, inserting at end")
        # 在末尾插入
        doc.add_page_break()

        heading = doc.add_heading('記憶體需求趨勢分析', level=3)

        caption = doc.add_paragraph('圖表 1：記憶體需求趨勢（近一年實測 + 未來五年預測）')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.bold = True
            run.font.size = Pt(10)

        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(str(chart_path), width=Inches(6.5))

    else:
        # 在找到的位置插入
        target_para = doc.paragraphs[best_position]
        parent = target_para._element.getparent()
        insert_idx = list(parent).index(target_para._element)

        # 創建標題段落
        heading = doc.add_heading('記憶體需求趨勢分析', level=3)
        heading_element = heading._element
        parent.insert(insert_idx, heading_element)

        # 創建圖片說明
        caption = doc.add_paragraph('圖表 1：記憶體需求趨勢（近一年實測 + 未來五年預測）')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.bold = True
            run.font.size = Pt(10)
        caption_element = caption._element
        parent.insert(insert_idx + 1, caption_element)

        # 創建圖片段落
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(str(chart_path), width=Inches(6.5))
        pic_element = pic_para._element
        parent.insert(insert_idx + 2, pic_element)

        # 空行
        space = doc.add_paragraph()
        space_element = space._element
        parent.insert(insert_idx + 3, space_element)

        print(f"[OK] Chart inserted at position {best_position}")

    # 保存
    doc.save(output_path)
    print(f"[OK] Document updated: {output_path.name}")

if __name__ == "__main__":
    doc_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")
    chart_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\memory_forecast_chart.png")
    output_path = doc_path

    print("Inserting memory forecast chart...")
    print(f"Document: {doc_path.name}")
    print(f"Chart: {chart_path.name}")
    print("-" * 80)

    insert_chart_to_docx(doc_path, chart_path, output_path)

    print("-" * 80)
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")
