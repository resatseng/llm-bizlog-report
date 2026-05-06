"""
将主管版 Markdown 转换为紧凑的 Word 文档（移除多余空白）
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from pathlib import Path

def parse_table(lines):
    """解析 Markdown 表格"""
    if not lines or len(lines) < 2:
        return None
    table_lines = [line for line in lines if not re.match(r'^\|[\s\-:]+\|$', line)]
    if len(table_lines) < 2:
        return None
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    return rows

def add_table_to_doc(doc, table_data):
    """添加表格到文档"""
    if not table_data:
        return
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                if i == 0:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                else:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                # 移除段落间距
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

def convert_md_to_docx_compact(input_file, output_file):
    """转换为紧凑的 Word 文档"""
    doc = Document()

    # 设定样式
    doc.styles['Normal'].font.name = '微軟正黑體'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(3)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.15

    # 标题样式
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        style.font.name = '微軟正黑體'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        style.font.color.rgb = RGBColor(0, 70, 140)
        style.paragraph_format.space_before = Pt(12 if i == 1 else 8)
        style.paragraph_format.space_after = Pt(6)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    skip_next_blank = False

    while i < len(lines):
        line = lines[i].rstrip()

        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(60, 60, 60)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
                skip_next_blank = True
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 处理表格
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            if i < len(lines) and not lines[i].strip().startswith('|'):
                table_data = parse_table(table_lines)
                if table_data:
                    add_table_to_doc(doc, table_data)
                in_table = False
                table_lines = []
                skip_next_blank = True
            continue

        # 空行 - 跳过多余的
        if not line.strip():
            if skip_next_blank:
                skip_next_blank = False
                i += 1
                continue
            # 不添加空段落，只设置标志
            i += 1
            continue

        # 分隔线 - 跳过
        if line.strip() == '---':
            i += 1
            continue

        # 标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            title_text = line.lstrip('#').strip()
            heading = doc.add_heading(title_text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            skip_next_blank = True
            i += 1
            continue

        # 列表
        if line.strip().startswith(('-', '*')):
            list_text = re.sub(r'^[\-\*]\s*', '', line.strip())
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
            p = doc.add_paragraph(list_text, style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # 数字列表
        if re.match(r'^\d+\.', line.strip()):
            list_text = re.sub(r'^\d+\.\s*', '', line.strip())
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
            p = doc.add_paragraph(list_text, style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # 一般段落
        text = line.strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)

        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                part = re.sub(r'`([^`]+)`', r'\1', part)
                p.add_run(part)

        i += 1

    doc.save(output_file)
    print(f"[OK] Compact document created")

if __name__ == "__main__":
    input_file = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃_主管版.md")
    output_file = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")

    print("Creating compact Word document...")
    print(f"Source: {input_file.name}")
    print(f"Output: {output_file.name}")
    print("-" * 80)

    convert_md_to_docx_compact(input_file, output_file)

    print("-" * 80)
    print(f"File size: {output_file.stat().st_size / 1024:.1f} KB")
