# -*- coding: utf-8 -*-
"""在規劃書 docx 末尾追加 v4.6 更新摘要"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pathlib

DOCX = pathlib.Path(r"d:\yujui\痛點需求地圖\prompt定版\業務日誌LLM萃取系統完整規劃書_v4_6.docx")
doc = Document(str(DOCX))

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 - (level - 1) * 2)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row

# ── 分頁符 ──────────────────────────────────────
doc.add_page_break()

# ── 標題 ────────────────────────────────────────
h = add_heading(doc, "更新摘要 v4.6（2026-04-24）", level=1)

add_para(doc, "本次更新涵蓋：Phase M/1 完成確認、Phase 2 執行狀態、商機等級架構釐清、痛需熱圖設計。")

# ── 進度表 ──────────────────────────────────────
add_heading(doc, "一、各 Phase 最新進度", level=2)
tbl = doc.add_table(rows=1, cols=4)
# tbl.style = "Table Grid"  # 此文件無表格樣式定義，略過
hdr = tbl.rows[0].cells
hdr[0].text = "Phase"
hdr[1].text = "狀態"
hdr[2].text = "數量"
hdr[3].text = "說明"

rows = [
    ("Phase 0", "✅ 完成", "82,105 筆",     "問卷三大信號"),
    ("Phase L", "✅ 完成", "1,802,590 筆",  "Step 0/1/2/3 零人工分類"),
    ("Phase 1", "✅ 完成", "206,817 家",    "8 大類法人屬性"),
    ("Phase M", "✅ 完成", "7 clusters",   "KMeans+PCA+熱路徑"),
    ("商機等級", "🔄 待全量", "756,989 筆",  "E/D/C2/C1/B/A 判定"),
    ("Phase 2", "🔄 6.8%", "~178,790 家",  "L1–L7 深度標籤（並行中）"),
    ("Phase 3", "⬜ 待建立", "—",           "痛需熱圖 + 三輸出"),
]
for r in rows:
    add_table_row(tbl, r)

# ── 架構釐清 ──────────────────────────────────────
add_heading(doc, "二、痛需熱圖架構釐清", level=2)
add_para(doc, "痛需熱圖由兩個獨立維度合成：")

tbl2 = doc.add_table(rows=1, cols=3)
# tbl2.style = "Table Grid"
h2 = tbl2.rows[0].cells
h2[0].text = "維度"
h2[1].text = "工具"
h2[2].text = "回答的問題"
for r in [
    ("L1–L7", "Phase 2 深度標籤", "客戶在討論什麼（痛點/角色/目標…）"),
    ("商機等級", "商機等級.ipynb", "這筆 Lead 走到哪一步（E→D→C2→C1→B→A）"),
]:
    add_table_row(tbl2, r)

add_para(doc, "熱圖 = L1 痛點（Y 軸）× 法人類型（X 軸），顏色深淺 = B+A 高階商機比例。")

# ── 並行執行說明 ──────────────────────────────────
add_heading(doc, "三、當前並行執行任務", level=2)
add_para(doc, "▶ Phase 2 深度標籤（執行中）：Gemini 2.0 Flash × 6 Workers，RESUME=True，6.8% 進度。")
add_para(doc, "▶ 商機等級全量（建議並行）：商機等級.ipynb 主流程已建立（斷點續傳+CSV輸出），可同步啟動。")
add_para(doc, "兩者資料來源獨立（Phase 2 用業務日報；商機等級用 CRMGY LeadInfo），互不干擾。")

# ── 儲存 ──────────────────────────────────────────
doc.save(str(DOCX))
print(f"✅ docx 已更新：{DOCX}")
