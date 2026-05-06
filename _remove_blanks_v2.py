"""
更激进地移除 Word 文档中的空白段落
"""
from docx import Document
from pathlib import Path

def remove_extra_blanks_v2(input_file, output_file):
    """移除多余的空白段落（更激进）"""
    doc = Document(input_file)

    original_count = len(doc.paragraphs)
    removed_count = 0

    # 标记要删除的段落索引
    to_remove = set()

    for i in range(len(doc.paragraphs)):
        p = doc.paragraphs[i]

        # 完全空白的段落
        if not p.text.strip():
            # 情况1：文档开头的空段落（保留0个）
            if i < 5:  # 前5个段落如果是空的，全部删除
                to_remove.add(i)

            # 情况2：连续的空段落（只保留1个）
            elif i > 0 and i in range(len(doc.paragraphs) - 1):
                prev_empty = not doc.paragraphs[i-1].text.strip()
                if prev_empty:
                    to_remove.add(i)

            # 情况3：标题后的空段落（删除）
            elif i > 0:
                prev_p = doc.paragraphs[i-1]
                # 检查上一段是否是标题
                if prev_p.style.name.startswith('Heading'):
                    to_remove.add(i)

    # 按倒序删除（避免索引变化）
    for idx in sorted(to_remove, reverse=True):
        p = doc.paragraphs[idx]
        p_element = p._element
        p_element.getparent().remove(p_element)
        removed_count += 1

    # 调整段落间距
    for p in doc.paragraphs:
        # 减小段落前后间距
        p_format = p.paragraph_format
        if p_format.space_before:
            p_format.space_before = min(p_format.space_before, 152400)  # 最多0.15英寸
        if p_format.space_after:
            p_format.space_after = min(p_format.space_after, 152400)

    doc.save(output_file)

    print(f"[OK] Optimization complete")
    print(f"  Original paragraphs: {original_count}")
    print(f"  Removed paragraphs: {removed_count}")
    print(f"  Final paragraphs: {original_count - removed_count}")
    if original_count > 0:
        print(f"  Reduction: {removed_count / original_count * 100:.1f}%")

if __name__ == "__main__":
    input_file = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")
    output_file = input_file

    print("Removing extra blank paragraphs (aggressive mode)...")
    print(f"File: {input_file.name}")
    print("-" * 80)

    remove_extra_blanks_v2(input_file, output_file)

    print("-" * 80)
    print(f"File size: {output_file.stat().st_size / 1024:.1f} KB")
