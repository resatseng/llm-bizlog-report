# -*- coding: utf-8 -*-
"""在規劃書 docx 末尾追加 v4.7 更新摘要"""
from docx import Document
from docx.shared import Pt
import pathlib

DOCX = pathlib.Path(r"d:\yujui\痛點需求地圖\prompt定版\業務日誌LLM萃取系統完整規劃書_v4_6.docx")
doc = Document(str(DOCX))

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 - (level - 1) * 2)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row

doc.add_page_break()

add_heading(doc, "更新摘要 v4.7（2026-04-27）", level=1)
add_para(doc, "本次更新涵蓋：商機等級升級為三層 TF-IDF KNN Pipeline、Phase 2 進度 40.5%、互動式 REPORT.html 公開儀表板。")

# 進度表
add_heading(doc, "一、各 Phase 最新進度（v4.7）", level=2)
tbl = doc.add_table(rows=1, cols=4)
hdr = tbl.rows[0].cells
hdr[0].text = "Phase"
hdr[1].text = "狀態"
hdr[2].text = "數量"
hdr[3].text = "說明"
rows = [
    ("Phase 0", "✅ 完成",    "82,105 筆",      "問卷三大信號"),
    ("Phase L", "✅ 完成",    "1,802,590 筆",   "Step 0/1/2/3 零人工分類"),
    ("Phase 1", "✅ 完成",    "206,817 家",     "8 大類法人屬性"),
    ("Phase M", "✅ 完成",    "7 clusters",    "KMeans+PCA+熱路徑；REPORT.html 公開"),
    ("商機等級", "🔄 2.0%",   "756,989 筆",     "三層 pipeline：關鍵詞→TF-IDF KNN→Batch LLM"),
    ("Phase 2", "🔄 40.5%",  "72,420/~178,790","L1–L7 深度標籤（RESUME=True）"),
    ("Phase 3", "⬜ 待建立",  "—",              "痛需熱圖 + 三輸出"),
]
for r in rows:
    add_table_row(tbl, r)

# 商機等級升級說明
add_heading(doc, "二、商機等級 Pipeline 升級（v4.7）", level=2)
add_para(doc, "原架構（Embedding API）因 Vertex AI service account 不支援 embed_content 而失敗。")
add_para(doc, "新架構採三層 Pipeline，完全無需 Embedding API：")

tbl2 = doc.add_table(rows=1, cols=4)
h2 = tbl2.rows[0].cells
h2[0].text = "層級"
h2[1].text = "方法"
h2[2].text = "覆蓋率"
h2[3].text = "API 呼叫"
for r in [
    ("Layer 1 關鍵詞快篩", "_KW + _NONE_KW 規則匹配", "~70%", "$0"),
    ("Layer 2 TF-IDF KNN", "char_wb 2-3gram，30k 特徵，k=5，信心≥0.65", "~25%", "$0"),
    ("Layer 3 Batch LLM",  "Gemini 2.5 Flash，10 筆/次批次呼叫", "~5%",   "極低"),
]:
    add_table_row(tbl2, r)

add_para(doc, "KNN 種子庫：14,999 筆種子 → 7,122 筆訓練（每類最多 3,000 平衡）。")
add_para(doc, "輸出欄位：E/D/C2/C1/B/A（bool）、stage_group、top_stage、stage_reason。")

# 互動式儀表板
add_heading(doc, "三、互動式 REPORT.html（公開儀表板）", level=2)
add_para(doc, "網址：https://resatseng.github.io/llm-bizlog-report/")
add_para(doc, "功能：")
add_para(doc, "  • PCA 散點圖：可依法人類型群組（C0–C6）篩選，支援全選 / 全消")
add_para(doc, "  • 雷達圖：7 維特徵，含單位（%、人、國、個）")
add_para(doc, "  • Lightbox：點擊縮圖放大雷達圖")

doc.save(str(DOCX))
print(f"OK: {DOCX}")
