"""
將伺服器排程資源計畫書.md 轉換為 Word 文檔（推薦系統規劃.docx）
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

def parse_table(lines):
    """解析 Markdown 表格"""
    if not lines or len(lines) < 2:
        return None

    # 移除分隔線
    table_lines = [line for line in lines if not re.match(r'^\|[\s\-:]+\|$', line)]

    if len(table_lines) < 2:
        return None

    # 解析表頭和數據
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)

    return rows

def add_table_to_doc(doc, table_data):
    """添加表格到文檔"""
    if not table_data:
        return

    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data

            # 表頭加粗
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

def convert_md_to_docx(input_file, output_file):
    """轉換 Markdown 為 Word 文檔"""
    doc = Document()

    # 設定中文字體
    doc.styles['Normal'].font.name = '微軟正黑體'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    doc.styles['Normal'].font.size = Pt(11)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # 處理代碼塊
        if line.startswith('```'):
            if in_code_block:
                # 結束代碼塊
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                code_lines = []
                in_code_block = False
            else:
                # 開始代碼塊
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 處理表格
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # 檢查下一行是否還是表格
            if i < len(lines) and not lines[i].strip().startswith('|'):
                # 表格結束
                table_data = parse_table(table_lines)
                if table_data:
                    add_table_to_doc(doc, table_data)
                    doc.add_paragraph()  # 空行
                in_table = False
                table_lines = []
            continue

        # 空行
        if not line.strip():
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue

        # 分隔線
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue

        # 標題
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            title_text = line.lstrip('#').strip()

            # 特殊處理：將文件標題改為「推薦系統規劃」
            if level == 1 and i == 0:
                title_text = '推薦系統規劃'

            heading = doc.add_heading(title_text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        # 列表
        if line.strip().startswith(('-', '*', '•')):
            list_text = re.sub(r'^[\-\*•]\s*', '', line.strip())
            # 處理 checkbox
            list_text = re.sub(r'^\[\s*\]\s*', '[ ] ', list_text)
            list_text = re.sub(r'^\[x\]\s*', '[x] ', list_text)

            p = doc.add_paragraph(list_text, style='List Bullet')
            i += 1
            continue

        # 數字列表
        if re.match(r'^\d+\.', line.strip()):
            list_text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue

        # 一般段落
        text = line.strip()

        # 處理粗體（**文字** 或 __文字__）
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                # 移除其他 markdown 語法
                part = re.sub(r'`([^`]+)`', r'\1', part)  # 內聯代碼
                p.add_run(part)

        i += 1

    # 保存文檔
    doc.save(output_file)
    print(f"[OK] 成功轉換：{output_file}")
    print(f"  標題已改為：推薦系統規劃")
    print(f"  版本：v2.1")

if __name__ == "__main__":
    input_file = Path(r"d:\yujui\痛點需求地圖\prompt定版\伺服器排程資源計畫書.md")
    output_file = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")

    print("開始轉換 Markdown 為 Word 文檔...")
    print(f"來源檔案：{input_file}")
    print(f"目標檔案：{output_file}")
    print("-" * 80)

    convert_md_to_docx(input_file, output_file)

    print("-" * 80)
    print(f"檔案大小：{output_file.stat().st_size / 1024:.1f} KB")
