"""
移除 Word 文档中的多余空白段落
"""
from docx import Document
from pathlib import Path

def remove_extra_blanks(input_file, output_file):
    """移除多余的空白段落"""
    doc = Document(input_file)

    # 统计
    original_paragraphs = len(doc.paragraphs)
    removed_count = 0

    # 反向遍历以安全删除
    paragraphs_to_remove = []

    for i in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[i]

        # 检查是否为空白段落（无文字且无表格）
        if not p.text.strip():
            # 检查前后是否都是空白段落（避免删除必要的分段）
            is_redundant = False

            # 如果是连续的空白段落，只保留一个
            if i > 0 and i < len(doc.paragraphs) - 1:
                prev_empty = not doc.paragraphs[i-1].text.strip()
                next_empty = not doc.paragraphs[i+1].text.strip()

                if prev_empty or next_empty:
                    is_redundant = True

            # 文档开头的空白段落
            elif i == 0:
                is_redundant = True

            # 文档结尾连续的空白段落（保留一个）
            elif i == len(doc.paragraphs) - 1 and i > 0:
                if not doc.paragraphs[i-1].text.strip():
                    is_redundant = True

            if is_redundant:
                paragraphs_to_remove.append(i)

    # 删除段落
    for idx in paragraphs_to_remove:
        p = doc.paragraphs[idx]
        p_element = p._element
        p_element.getparent().remove(p_element)
        removed_count += 1

    # 保存
    doc.save(output_file)

    final_paragraphs = original_paragraphs - removed_count

    print(f"[OK] Optimization complete")
    print(f"  Original paragraphs: {original_paragraphs}")
    print(f"  Removed paragraphs: {removed_count}")
    print(f"  Final paragraphs: {final_paragraphs}")
    if original_paragraphs > 0:
        print(f"  Optimization rate: {removed_count / original_paragraphs * 100:.1f}%")

if __name__ == "__main__":
    input_file = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")
    output_file = input_file  # 覆盖原文件

    print("Start optimizing Word document...")
    print(f"File: {input_file}")
    print("-" * 80)

    remove_extra_blanks(input_file, output_file)

    print("-" * 80)
    print(f"File size: {output_file.stat().st_size / 1024:.1f} KB")
