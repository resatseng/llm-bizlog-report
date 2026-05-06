"""
將記憶體需求圖表插入 Word 文檔
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def insert_chart_after_section(doc_path, chart_path, output_path, search_text):
    """在指定章節後插入圖表"""
    doc = Document(doc_path)
    chart_inserted = False

    # 尋找目標段落（在「長期成本演進摘要」之後）
    for i, paragraph in enumerate(doc.paragraphs):
        if search_text in paragraph.text:
            # 在這個段落後面插入圖表
            # 先添加標題
            new_heading_element = paragraph._element
            parent = new_heading_element.getparent()

            # 找到插入位置
            insert_idx = list(parent).index(new_heading_element) + 1

            # 創建新段落（圖表標題）
            caption_para = doc.add_paragraph()
            caption_para.text = '圖表 1：記憶體需求趨勢分析（近一年實測 + 未來五年預測）'
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_para.runs:
                run.font.bold = True
                run.font.size = Pt(11)

            # 移動到正確位置
            caption_element = caption_para._element
            parent.insert(insert_idx, caption_element)

            # 創建圖片段落
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 插入圖片（寬度6.5英寸，保持比例）
            run = pic_para.add_run()
            run.add_picture(str(chart_path), width=Inches(6.5))

            # 移動圖片到正確位置
            pic_element = pic_para._element
            parent.insert(insert_idx + 1, pic_element)

            # 添加空行
            space_para = doc.add_paragraph()
            space_element = space_para._element
            parent.insert(insert_idx + 2, space_element)

            chart_inserted = True
            print(f"[OK] Chart inserted after: {search_text[:50]}...")
            break

    if not chart_inserted:
        print(f"[WARNING] Target section not found: {search_text}")
        # 如果找不到，就插入到文檔末尾
        doc.add_page_break()
        caption = doc.add_paragraph('圖表 1：記憶體需求趨勢分析（近一年實測 + 未來五年預測）')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.bold = True
            run.font.size = Pt(11)

        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(str(chart_path), width=Inches(6.5))
        print("[OK] Chart inserted at end of document")

    # 保存
    doc.save(output_path)
    print(f"[OK] Document saved: {output_path.name}")

if __name__ == "__main__":
    doc_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")
    chart_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\memory_forecast_chart.png")
    output_path = doc_path  # 覆蓋原文件

    print("Inserting chart into Word document...")
    print(f"Document: {doc_path.name}")
    print(f"Chart: {chart_path.name}")
    print("-" * 80)

    # 在「長期成本演進摘要」之後插入圖表
    insert_chart_after_section(
        doc_path,
        chart_path,
        output_path,
        search_text="長期成本演進摘要"
    )

    print("-" * 80)
    print(f"Final size: {output_path.stat().st_size / 1024:.1f} KB")
