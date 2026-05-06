"""
插入兩個圖表到 Word 文檔：
1. 記憶體需求趨勢圖
2. SQL 資源分析圖
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def find_section_position(doc, keywords):
    """尋找包含關鍵詞的段落位置"""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(keyword in text for keyword in keywords):
            # 向後找 5-10 個段落，找到下一個標題
            for j in range(i+1, min(i+15, len(doc.paragraphs))):
                if doc.paragraphs[j].style.name.startswith('Heading'):
                    return j
    return None

def insert_chart(doc, position, chart_path, title, chart_num):
    """在指定位置插入圖表"""
    if position is None:
        # 在末尾插入
        doc.add_page_break()
        heading = doc.add_heading(title.split('：')[0], level=3)

        caption = doc.add_paragraph(f'{title}')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.bold = True
            run.font.size = Pt(10)

        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(str(chart_path), width=Inches(6.5))

        print(f"[OK] Chart {chart_num} inserted at end")
    else:
        # 在指定位置插入
        target_para = doc.paragraphs[position]
        parent = target_para._element.getparent()
        insert_idx = list(parent).index(target_para._element)

        # 標題
        heading = doc.add_heading(title.split('：')[0], level=3)
        heading_element = heading._element
        parent.insert(insert_idx, heading_element)

        # 圖片說明
        caption = doc.add_paragraph(f'{title}')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.bold = True
            run.font.size = Pt(10)
        caption_element = caption._element
        parent.insert(insert_idx + 1, caption_element)

        # 圖片
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(str(chart_path), width=Inches(6.5))
        pic_element = pic_para._element
        parent.insert(insert_idx + 2, pic_element)

        # 空行
        space = doc.add_paragraph()
        space_element = space._element
        parent.insert(insert_idx + 3, space_element)

        print(f"[OK] Chart {chart_num} inserted at position {position}")

def insert_both_charts(doc_path, output_path):
    """插入兩個圖表"""
    doc = Document(doc_path)

    # 圖表 1：記憶體需求趨勢（在「長期成本演進摘要」之後）
    memory_chart_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\memory_forecast_chart.png")
    pos1 = find_section_position(doc, ["長期成本", "5年總成本", "基準情境"])

    insert_chart(
        doc, pos1,
        memory_chart_path,
        "圖表 1：記憶體需求趨勢（近一年實測 + 未來五年預測）",
        1
    )

    # 重新載入文檔（因為插入後段落索引會變）
    doc.save(output_path)
    doc = Document(output_path)

    # 圖表 2：SQL 資源分析（在「SQL Server 查詢資源評估」之後）
    sql_chart_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\sql_resource_analysis.png")
    pos2 = find_section_position(doc, ["SQL Server", "資源消耗", "查詢資源評估"])

    insert_chart(
        doc, pos2,
        sql_chart_path,
        "圖表 2：SQL Server 資源消耗與效能分析",
        2
    )

    # 保存
    doc.save(output_path)
    print(f"\n[OK] Document updated with 2 charts: {output_path.name}")

if __name__ == "__main__":
    doc_path = Path(r"d:\yujui\痛點需求地圖\prompt定版\推薦系統規劃.docx")
    output_path = doc_path

    print("Inserting both charts into document...")
    print(f"Document: {doc_path.name}")
    print("-" * 80)

    insert_both_charts(doc_path, output_path)

    print("-" * 80)
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")
